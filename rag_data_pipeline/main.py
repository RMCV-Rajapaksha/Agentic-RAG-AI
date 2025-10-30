# Core LlamaIndex imports
from llama_index.core import Document
from llama_index.core.ingestion import IngestionPipeline
from llama_index.core.node_parser import MarkdownNodeParser
from llama_index.core.extractors import TitleExtractor
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.core.embeddings import BaseEmbedding
from pydantic import PrivateAttr
from azure.ai.inference import EmbeddingsClient
from azure.core.credentials import AzureKeyCredential

# Local imports
from database.db import DatabaseConnection
from src.youtube_transcripts.youtube_transcript_to_md import get_transcript_segments
from src.scraper.web_scraper import get_markdown
from src.drive_reader.drive_reader import load_google_drive_documents
from config.config import get_azure_endpoint_embedding, get_azure_api_key_embedding, get_google_drive_folder_id

# Standard imports
import os
from typing import List
from pathlib import Path

# External lightweight libraries
import pypandoc
import pdfplumber
from docx import Document as DocxDocument
from markdownify import markdownify as md
import requests
import re

AZURE_ENDPOINT = get_azure_endpoint_embedding()
AZURE_API_KEY = get_azure_api_key_embedding()
DEPLOYMENT_NAME = "text-embedding-3-small"


class AzureAIEmbedding(BaseEmbedding):
    """Wrapper for Azure AI Foundry embeddings to work with llama_index"""
    
    _endpoint: str = PrivateAttr()
    _api_key: str = PrivateAttr()
    _deployment: str = PrivateAttr()
    _embed_dim: int = PrivateAttr(default=1536)
    _client: EmbeddingsClient = PrivateAttr()
    
    def __init__(self, endpoint: str, api_key: str, deployment: str, embed_dim: int = 1536, **kwargs):
        super().__init__(**kwargs)
        self._endpoint = endpoint
        self._api_key = api_key
        self._deployment = deployment
        self._embed_dim = embed_dim
        self._client = EmbeddingsClient(
            endpoint=f"{endpoint}openai/deployments/{deployment}",
            credential=AzureKeyCredential(api_key)
        )
    
    @property
    def embed_dim(self) -> int:
        """Return the embedding dimension."""
        return self._embed_dim
    
    def _get_query_embedding(self, query: str) -> list[float]:
        """Get embedding for a query text (required by llama_index)"""
        response = self._client.embed(input=[query])
        return response.data[0].embedding
    
    def _get_text_embedding(self, text: str) -> list[float]:
        """Get embedding for a single text"""
        response = self._client.embed(input=[text])
        return response.data[0].embedding
    
    async def _aget_query_embedding(self, query: str) -> list[float]:
        """Async version of get_query_embedding"""
        return self._get_query_embedding(query)
    
    def _get_text_embeddings(self, texts: list[str]) -> list[list[float]]:
        """Get embeddings for multiple texts"""
        response = self._client.embed(input=texts)
        return [item.embedding for item in response.data]


class LightweightConverter:
    def convert(self, source: str) -> str:
        """Convert document to Markdown using lightweight libraries."""
        ext = Path(source).suffix.lower()

        try:
            if ext in [".docx", ".pptx", ".odt"]:
                # Use pypandoc for docx/pptx/odt
                return pypandoc.convert_file(source, "md", extra_args=["--wrap=none"])

            elif ext == ".pdf":
                # Extract text from PDF
                text = ""
                with pdfplumber.open(source) as pdf:
                    text = "\n".join([page.extract_text() or "" for page in pdf.pages])
                return text

            elif ext == ".txt":
                # Plain text file
                with open(source, "r", encoding="utf-8") as f:
                    return f.read()

            elif ext == ".html" or ext == ".htm":
                # Convert HTML → Markdown
                with open(source, "r", encoding="utf-8") as f:
                    html_content = f.read()
                return md(html_content)

            else:
                print(f"Unsupported format for {source}")
                return ""

        except Exception as e:
            print(f"Error converting {source}: {e}")
            return ""


# ===============================
# Main Ingestion Class
# ===============================
class RAGDataIngestion:
    """
    RAG data ingestion pipeline - ingests web URLs, Google Drive docs, and YouTube transcripts into DB
    """

    def __init__(self):
        self.db_connection = DatabaseConnection()

        self.document_converter = LightweightConverter
        self.document_converter = LightweightConverter()

        self.vector_store = self.db_connection.get_vector_store()

        self.pipeline = IngestionPipeline(
            transformations=[
                MarkdownNodeParser(chunk_size=512, chunk_overlap=100, include_metadata=True, include_prev_next_rel=True),
                TitleExtractor(),
            
                AzureAIEmbedding(
                    endpoint=AZURE_ENDPOINT,
                    api_key=AZURE_API_KEY,
                    deployment=DEPLOYMENT_NAME,
                    embed_dim=1536
                )
            ],
            vector_store=self.vector_store,
        )

    def convert_document_to_markdown(self, source: str) -> str:
        return self.document_converter.convert(source)

    def scrape_web_urls(self, urls: List[str]) -> List[Document]:
        documents = []
        print("Scraping web URLs for markdown content...")
        for url in urls:
            scraped_data = get_markdown(url)
            if scraped_data:
                doc = Document(
                    text=scraped_data['content_markdown'],
                    metadata={
                        'url': scraped_data['url'],
                        'title': scraped_data['metadata']['title'],
                        'description': scraped_data['metadata']['description'],
                        'source': 'web_scraper'
                    }
                )
                documents.append(doc)
        return documents

    def load_drive_documents(self, folder_id: str) -> List[Document]:
        """Load documents from Google Drive using functional approach."""
        documents = load_google_drive_documents(folder_id)
        for doc in documents:
            doc.metadata['source'] = 'google_drive'
            doc.metadata['folder_id'] = folder_id
        return documents

    def convert_drive_documents_to_markdown(self, folder_id: str) -> List[Document]:
        """Load and convert Google Drive documents to markdown using functional approach."""
        documents = []
        drive_documents = load_google_drive_documents(folder_id)

        for doc in drive_documents:
            if 'file_path' in doc.metadata:
                file_path = doc.metadata['file_path']
                print(f"Converting document from Google Drive: {file_path}")
                markdown_content = self.convert_document_to_markdown(file_path)

                if markdown_content:
                    converted_doc = Document(
                        text=markdown_content,
                        metadata={
                            'source': 'google_drive_converted',
                            'folder_id': folder_id,
                            'original_file_path': file_path,
                            'type': 'converted_document'
                        }
                    )
                    documents.append(converted_doc)
            else:
                doc.metadata['source'] = 'google_drive'
                doc.metadata['folder_id'] = folder_id
                documents.append(doc)

        return documents

    def process_youtube_videos(self, urls: List[str], segment_length_minutes: int = 10) -> List[Document]:
        """Process YouTube videos and return document segments
        
        Args:
            urls: List of YouTube video URLs
            segment_length_minutes: Length of each segment in minutes (default: 10)
            
        Returns:
            List of Document objects, one per segment
        """
        documents = []
        print("Processing YouTube videos for transcript segments...")
        for link in urls:
            try:
                video_data = get_transcript_segments(link, language="en", segment_length_minutes=segment_length_minutes)
                for segment in video_data['segments']:
                    video_doc = Document(
                        text=segment['content_markdown'],
                        metadata={
                            'url': video_data['url'],
                            'title': video_data['metadata'].get('title', ''),
                            'description': video_data['metadata'].get('description', ''),
                            'source': 'youtube_transcript',
                            'start_seconds': segment['start_seconds'],
                            'end_seconds': segment['end_seconds'],
                        }
                    )
                    documents.append(video_doc)
                print(f"Processed {len(video_data['segments'])} segments from {link}")
            except Exception as e:
                print(f"Error processing YouTube video {link}: {e}")
        return documents

    def ingest_documents(self, documents: List[Document]) -> None:
        if not documents:
            return

        filtered_documents = []
        try:
            existing_urls = set()
            try:
                all_nodes = self.vector_store._get_all() if hasattr(self.vector_store, '_get_all') else []
                for node in all_nodes:
                    url = node.metadata.get('url') if hasattr(node, 'metadata') and node.metadata else None
                    if url:
                        existing_urls.add(url)
            except Exception as e:
                print(f"Could not fetch existing URLs from vector store: {e}")

            for doc in documents:
                url = doc.metadata.get('url') if doc.metadata else None
                if url and url in existing_urls:
                    print(f"Skipping duplicate URL: {url}")
                    continue
                filtered_documents.append(doc)
        except Exception as e:
            print(f"Error during duplicate filtering: {e}")
            filtered_documents = documents

        if not filtered_documents:
            print("No new documents to ingest after filtering.")
            return

        self.pipeline.run(documents=filtered_documents, show_progress=True)


# ===============================
# Main Entry Point
# ===============================
def main():
    pipeline = RAGDataIngestion()

    urls_to_scrape = [
        # "https://wso2.ai/",
        # "https://wso2.com/api-management/ai/",
        # "https://wso2.com/integration/ai/",
        # "https://wso2.com/identity-and-access-management/ai/",
        # "https://wso2.com/internal-developer-platform/ai/"
    ]

    # Fetch YouTube URLs from GitHub markdown file
    github_md_url = "https://raw.githubusercontent.com/RMCV-Rajapaksha/Agentic-RAG-AI/main/YouTubeURL.md"
    
    try:
        response = requests.get(github_md_url)
        response.raise_for_status()
        md_content = response.text
        
        # Extract YouTube URLs using regex
        youtube_url_pattern = r'https://www\.youtube\.com/watch\?v=[\w-]+'
        urls_to_videos = re.findall(youtube_url_pattern, md_content)
        
        print(f"Found {len(urls_to_videos)} YouTube URLs from markdown file")
    except Exception as e:
        print(f"Error fetching URLs from markdown: {e}")
        urls_to_videos = []
    
    print(f"YouTube URLs to process: {urls_to_videos}")


    drive_folder_id = get_google_drive_folder_id()

    try:
        all_documents = []

        if urls_to_videos:
            youtube_documents = pipeline.process_youtube_videos(urls_to_videos)
            all_documents.extend(youtube_documents)

        if urls_to_scrape:
            url_documents = pipeline.scrape_web_urls(urls_to_scrape)
            all_documents.extend(url_documents)

        if drive_folder_id:
            print("Loading and converting Google Drive documents...")
            drive_documents = pipeline.convert_drive_documents_to_markdown(drive_folder_id)
            print(f"Loaded and converted {len(drive_documents)} documents from Google Drive.")
            all_documents.extend(drive_documents)

        if all_documents:
            print(f"Ingesting {len(all_documents)} documents...")
            pipeline.ingest_documents(all_documents)
            print("✅ Data ingestion completed successfully!")
        else:
            print("⚠️ No documents to ingest.")

    except Exception as e:
        print(f"❌ An error occurred during ingestion: {e}")


if __name__ == "__main__":
    main()
